#!/usr/bin/env python3
"""Build a polished Chinese DOCX from checkpointed translation JSON.

Input schema:
{
  "title_zh": "中文标题",
  "subtitle_zh": "可选副标题",
  "source_note": "可选译者说明",
  "pages": [
    {
      "page": 1,
      "translated_paragraphs": [
        {"type": "heading1", "text": "1 标题"},
        {"type": "heading2", "text": "1.1 小标题"},
        {"type": "paragraph", "text": "正文"},
        {"type": "bullet", "text": "项目"},
        {"type": "number", "text": "步骤"},
        {
          "type": "table",
          "headers": ["字段", "内容"],
          "rows": [["项目名称", "示例"]]
        }
      ]
    }
  ]
}

For convenience, translated_paragraphs may also contain plain strings.
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLUE = "1769AA"
LIGHT_BLUE = "E8F1F8"
GRAY = "666666"
FONT = "PingFang SC"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run(str(text)), 9, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, 8, color=GRAY)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.28

    for name, size, before, after in [
        ("Title", 26, 0, 10),
        ("Heading 1", 17, 16, 8),
        ("Heading 2", 13.5, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_text(doc, item):
    if isinstance(item, str):
        item = {"type": "paragraph", "text": item}
    kind = item.get("type", "paragraph")
    text = item.get("text", "").strip()
    if kind == "table":
        headers = item.get("headers", [])
        rows = item.get("rows", [])
        column_count = max(len(headers), max((len(row) for row in rows), default=0))
        if not column_count:
            return
        table = doc.add_table(rows=1 if headers else 0, cols=column_count)
        table.autofit = True
        if headers:
            for index, value in enumerate(headers):
                cell_text(table.cell(0, index), value, True, "FFFFFF")
                shade(table.cell(0, index), BLUE)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cell_text(cells[index], value)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return
    if not text:
        return
    styles = {
        "heading1": "Heading 1",
        "heading2": "Heading 2",
        "heading3": "Heading 3",
        "bullet": "List Bullet",
        "number": "List Number",
    }
    paragraph = doc.add_paragraph(style=styles.get(kind))
    paragraph.paragraph_format.widow_control = True
    set_font(paragraph.add_run(text), 10.5)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("translation_json")
    parser.add_argument("output_docx")
    args = parser.parse_args()

    source = Path(args.translation_json)
    output = Path(args.output_docx)
    payload = json.loads(source.read_text(encoding="utf-8"))

    doc = Document()
    configure(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(payload.get("title_zh", "中文译本")), 8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_number(footer)

    title = doc.add_paragraph(style="Title")
    set_font(title.add_run(payload.get("title_zh", "中文译本")), 26, True, BLUE)
    if payload.get("subtitle_zh"):
        subtitle = doc.add_paragraph()
        set_font(subtitle.add_run(payload["subtitle_zh"]), 13, color=GRAY)
    if payload.get("source_note"):
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(16)
        set_font(note.add_run(payload["source_note"]), 9, color=GRAY)
    doc.add_page_break()

    for page in payload.get("pages", []):
        for item in page.get("translated_paragraphs", []):
            add_text(doc, item)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = payload.get("title_zh", "中文译本")
    doc.core_properties.subject = "PDF Chinese Translation"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
